import streamlit as st
import os
import re
from docx import Document
from io import StringIO
from io import BytesIO
import qrcode
import base64

def extract_dois_from_text(text):
    # Regular expression to find DOIs starting with "10."
    doi_pattern = r'10\.\d{4,9}/[^\s\)\]\[]+(?:\([^\s\)\]\[]+\))?'
    matches = re.findall(doi_pattern, text)

    # Clean up DOIs by removing extraneous characters
    cleaned_matches = []
    for match in matches:
        match = re.sub(r'[\[\]\)]', '', match)  # Remove '[', ']', ')'
        match = match.strip('.')
        cleaned_matches.append(match)
    
    return cleaned_matches

def extract_dois_from_paragraphs(paragraphs):
    full_text = []
    for para in paragraphs:
        full_text.append(para.text)
    text = '\n'.join(full_text)
    return extract_dois_from_text(text)

def extract_dois_from_tables(tables):
    full_text = []
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                full_text.append(cell.text)
    text = '\n'.join(full_text)
    return extract_dois_from_text(text)

def extract_dois_from_word(docx_file):
    # Load the Word document
    doc = Document(docx_file)
    
    # Extract DOIs from paragraphs
    dois = extract_dois_from_paragraphs(doc.paragraphs)
    
    # Extract DOIs from tables
    dois += extract_dois_from_tables(doc.tables)
    
    return dois

def save_dois_to_text(dois):
    # Create a StringIO object to store DOIs
    output = StringIO()
    for doi in sorted(set(dois)):  # Remove duplicates and sort
        output.write(doi.strip() + '\n')
    return output.getvalue()

# Streamlit app layout
st.title("DOI Extractor from Word Documents")
st.write("Upload multiple .docx files to extract DOIs. The extracted DOIs will be saved in a text file for download.")

# File uploader for multiple files
uploaded_files = st.file_uploader("Upload .docx files", type='docx', accept_multiple_files=True)

if uploaded_files:
    all_dois = set()  # Use a set to collect DOIs and automatically remove duplicates
    
    for uploaded_file in uploaded_files:
        # Read the uploaded Word file
        st.write(f"Processing file: {uploaded_file.name}")
        dois = extract_dois_from_word(uploaded_file)
        all_dois.update(dois)

    # Save the extracted DOIs to a text string
    extracted_dois = save_dois_to_text(all_dois)

    # Prepare the text for download
    if extracted_dois:
        st.download_button(
            label="Download Extracted DOIs",
            data=extracted_dois,
            file_name='extracted_dois.txt',
            mime='text/plain'
        )
        st.success(f"Extracted {len(all_dois)} unique DOIs.")
    else:
        st.warning("No DOIs were extracted from the uploaded files.")
else:
    st.info("Please upload .docx files to extract DOIs.")

st.info("Created by Dr. Satyajeet Patil")
st.info("For more cool apps like this visit: https://patilsatyajeet.wixsite.com/home/python")

# Title of the section
st.title("Support our Research")
st.write("Scan the QR code below to make a payment to: satyajeet1396@oksbi")

# Generate the UPI QR code
upi_url = "upi://pay?pa=satyajeet1396@oksbi&pn=Satyajeet Patil&cu=INR"
qr = qrcode.make(upi_url)

# Save the QR code image to a BytesIO object
buffer = BytesIO()
qr.save(buffer, format="PNG")
buffer.seek(0)

# Convert the image to Base64
qr_base64 = base64.b64encode(buffer.getvalue()).decode()

# Center-align the QR code image using HTML and CSS
st.markdown(
    f"""
    <div style="display: flex; justify-content: center; align-items: center;">
        <img src="data:image/png;base64,{qr_base64}" width="200">
    </div>
    """,
    unsafe_allow_html=True
)

# Display the "Buy Me a Coffee" button as an image link
st.markdown(
    """
    <div style="text-align: center; margin-top: 20px;">
        <a href="https://www.buymeacoffee.com/researcher13" target="_blank">
            <img src="https://img.buymeacoffee.com/button-api/?text=Support our Research&emoji=&slug=researcher13&button_colour=FFDD00&font_colour=000000&font_family=Cookie&outline_colour=000000&coffee_colour=ffffff" alt="Support our Research"/>
        </a>
    </div>
    """,
    unsafe_allow_html=True
)
st.info("A small donation from you can fuel our research journey, turning ideas into breakthroughs that can change lives!")
